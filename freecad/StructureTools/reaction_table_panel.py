import FreeCAD, FreeCADGui
import os
from typing import List, Dict, Any, Optional
import csv
import json

# Prefer PySide2 when available
try:
    from PySide2 import QtWidgets, QtCore, QtGui
except ImportError:
    try:
        from PySide import QtWidgets, QtCore, QtGui
    except ImportError as e:
        raise ImportError("Neither PySide2 nor PySide could be imported. Please install one of them.") from e

ICONPATH = os.path.join(os.path.dirname(__file__), "resources")


class ReactionTablePanel:
    """Panel for displaying reaction results in a table format with export functionality."""
    
    def __init__(self, reaction_obj):
        self.reaction_obj = reaction_obj
        self.form = self.create_ui()
        self.setup_connections()
        self.populate_reaction_table()
    
    def create_ui(self):
        """Create the user interface for the reaction table panel."""
        # Main widget
        widget = QtWidgets.QWidget()
        main_layout = QtWidgets.QVBoxLayout(widget)
        
        # Title
        title_label = QtWidgets.QLabel("Reaction Results Table")
        title_label.setStyleSheet("font-weight: bold; font-size: 14px; color: #2c3e50;")
        main_layout.addWidget(title_label)
        
        # Information label
        info_text = "Reaction forces and moments at supported nodes for the selected load combination."
        self.info_label = QtWidgets.QLabel(info_text)
        self.info_label.setStyleSheet("color: #7f8c8d; font-style: italic;")
        main_layout.addWidget(self.info_label)
        
        # Load combination selection
        combo_layout = QtWidgets.QHBoxLayout()
        combo_layout.addWidget(QtWidgets.QLabel("Load Combination:"))
        self.load_combo_dropdown = QtWidgets.QComboBox()
        self.load_combo_dropdown.setMinimumWidth(150)
        combo_layout.addWidget(self.load_combo_dropdown)
        combo_layout.addStretch()
        main_layout.addLayout(combo_layout)
        
        # Table for reaction results
        self.table_widget = QtWidgets.QTableWidget()
        self.table_widget.setAlternatingRowColors(True)
        self.table_widget.setStyleSheet("""
            QTableWidget {
                alternate-background-color: #f0f0f0;
                background-color: white;
            }
            QTableWidget::item:selected {
                background-color: #3498db;
                color: white;
            }
        """)
        self.table_widget.setSortingEnabled(True)
        main_layout.addWidget(self.table_widget)
        
        # Control buttons
        button_layout = QtWidgets.QHBoxLayout()
        
        self.refresh_button = QtWidgets.QPushButton("Refresh")
        self.refresh_button.setIcon(QtGui.QIcon.fromTheme("view-refresh"))
        
        self.export_excel_button = QtWidgets.QPushButton("Export to Excel")
        self.export_excel_button.setIcon(QtGui.QIcon.fromTheme("document-export"))
        
        self.export_word_button = QtWidgets.QPushButton("Export to Word")
        self.export_word_button.setIcon(QtGui.QIcon.fromTheme("document-export"))
        
        self.export_csv_button = QtWidgets.QPushButton("Export to CSV")
        self.export_csv_button.setIcon(QtGui.QIcon.fromTheme("document-export"))
        
        button_layout.addWidget(self.refresh_button)
        button_layout.addStretch()
        button_layout.addWidget(self.export_excel_button)
        button_layout.addWidget(self.export_word_button)
        button_layout.addWidget(self.export_csv_button)
        
        main_layout.addLayout(button_layout)
        
        # Status bar
        self.status_label = QtWidgets.QLabel("Ready")
        self.status_label.setStyleSheet("color: #27ae60;")
        main_layout.addWidget(self.status_label)
        
        return widget
    
    def setup_connections(self):
        """Connect UI elements to their handlers."""
        self.load_combo_dropdown.currentTextChanged.connect(self.on_load_combination_changed)
        self.refresh_button.clicked.connect(self.populate_reaction_table)
        self.export_excel_button.clicked.connect(lambda: self.export_to_format("excel"))
        self.export_word_button.clicked.connect(lambda: self.export_to_format("word"))
        self.export_csv_button.clicked.connect(lambda: self.export_to_format("csv"))
    
    def populate_load_combinations(self):
        """Populate the load combination dropdown."""
        self.load_combo_dropdown.clear()
        
        if not self.reaction_obj or not self.reaction_obj.ObjectBaseCalc:
            return
        
        try:
            # Get available load combinations from calc object
            calc_obj = self.reaction_obj.ObjectBaseCalc
            model = None
            
            # Try different ways to get the FE model
            if hasattr(calc_obj, 'FEModel') and calc_obj.FEModel:
                model = calc_obj.FEModel
            elif hasattr(calc_obj, 'model') and calc_obj.model:
                model = calc_obj.model
            elif hasattr(calc_obj, 'Proxy') and hasattr(calc_obj.Proxy, 'model') and calc_obj.Proxy.model:
                model = calc_obj.Proxy.model
            
            # First try to get combinations from the model
            if model and hasattr(model, 'LoadCombos') and model.LoadCombos:
                combo_names = list(model.LoadCombos.keys())
                self.load_combo_dropdown.addItems(combo_names)
                
                # Set current selection to match object property
                if hasattr(self.reaction_obj, 'ActiveLoadCombination'):
                    current_combo = self.reaction_obj.ActiveLoadCombination
                    index = self.load_combo_dropdown.findText(current_combo)
                    if index >= 0:
                        self.load_combo_dropdown.setCurrentIndex(index)
                    elif combo_names:
                        # Set to first available combination
                        self.load_combo_dropdown.setCurrentIndex(0)
                        # Update object property
                        self.reaction_obj.ActiveLoadCombination = combo_names[0]
                return
            
            # Try to get combinations directly from calc object
            if hasattr(calc_obj, 'LoadCombination'):
                available_combos = calc_obj.LoadCombination
                
                if isinstance(available_combos, list):
                    self.load_combo_dropdown.addItems(available_combos)
                else:
                    self.load_combo_dropdown.addItem(str(available_combos))
                
                # Set current selection to match object property
                if hasattr(self.reaction_obj, 'ActiveLoadCombination'):
                    current_combo = self.reaction_obj.ActiveLoadCombination
                    index = self.load_combo_dropdown.findText(current_combo)
                    if index >= 0:
                        self.load_combo_dropdown.setCurrentIndex(index)
                return
            
            # If no combinations found, add default
            self.load_combo_dropdown.addItem("100_DL")
            if hasattr(self.reaction_obj, 'ActiveLoadCombination') and not self.reaction_obj.ActiveLoadCombination:
                self.reaction_obj.ActiveLoadCombination = "100_DL"                  
                
        except Exception as e:
            FreeCAD.Console.PrintWarning(f"Warning: Could not populate load combinations: {str(e)}\n")
            # Add default combination
            self.load_combo_dropdown.addItem("100_DL")
    
    def populate_reaction_table(self):
        """Populate the reaction table with data."""
        try:
            # Clear existing table
            self.table_widget.clear()
            
            if not self.reaction_obj or not self.reaction_obj.ObjectBaseCalc:
                self.status_label.setText("No calculation object linked")
                self.status_label.setStyleSheet("color: #e74c3c;")
                return
            
            # Populate load combinations
            self.populate_load_combinations()
            
            # Get current load combination
            load_combo = self.load_combo_dropdown.currentText()
            if not load_combo:
                self.status_label.setText("No load combinations available")
                self.status_label.setStyleSheet("color: #e74c3c;")
                return
            
            calc_obj = self.reaction_obj.ObjectBaseCalc
            model = None
            
            # Try different ways to get the FE model
            if hasattr(calc_obj, 'FEModel') and calc_obj.FEModel:
                model = calc_obj.FEModel
                FreeCAD.Console.PrintMessage("Using FEModel attribute\n")
            elif hasattr(calc_obj, 'model') and calc_obj.model:
                model = calc_obj.model
                FreeCAD.Console.PrintMessage("Using model attribute\n")
            elif hasattr(calc_obj, 'Proxy') and hasattr(calc_obj.Proxy, 'model') and calc_obj.Proxy.model:
                model = calc_obj.Proxy.model
                FreeCAD.Console.PrintMessage("Using Proxy.model attribute\n")
                
            if not model or not hasattr(model, 'nodes') or not model.nodes:
                self.status_label.setText("Analysis not completed or no nodes available - please run calculation")
                self.status_label.setStyleSheet("color: #e74c3c;")
                return
            
            # Prepare table data
            table_data = []
            headers = ["Node ID", "X (m)", "Y (m)", "Z (m)", "FX (kN)", "FY (kN)", "FZ (kN)", "MX (kN·m)", "MY (kN·m)", "MZ (kN·m)"]
            
            # Import units manager for formatting
            try:
                from .utils.units_manager import format_force, format_moment
                units_available = True
            except ImportError:
                units_available = False
                format_force = lambda x: f"{x:.3f}"
                format_moment = lambda x: f"{x:.3f}"
            
            # Prefer using stored reaction properties if they exist
            if (hasattr(calc_obj, 'ReactionNodes') and calc_obj.ReactionNodes and
                hasattr(calc_obj, 'ReactionX') and calc_obj.ReactionX and
                hasattr(calc_obj, 'ReactionY') and calc_obj.ReactionY and
                hasattr(calc_obj, 'ReactionZ') and calc_obj.ReactionZ and
                hasattr(calc_obj, 'ReactionMX') and calc_obj.ReactionMX and
                hasattr(calc_obj, 'ReactionMY') and calc_obj.ReactionMY and
                hasattr(calc_obj, 'ReactionMZ') and calc_obj.ReactionMZ):
                
                FreeCAD.Console.PrintMessage("Using stored reaction properties\n")
                
                # Collect reaction data for each supported node from properties
                for i, node_name in enumerate(calc_obj.ReactionNodes):
                    # Get node coordinates from model if possible
                    x_pos = y_pos = z_pos = 0.0
                    
                    if model and node_name in model.nodes:
                        x_pos = model.nodes[node_name].X
                        y_pos = model.nodes[node_name].Y
                        z_pos = model.nodes[node_name].Z
                    
                    # Get reaction values from stored properties
                    rx = calc_obj.ReactionX[i] if i < len(calc_obj.ReactionX) else 0.0
                    ry = calc_obj.ReactionY[i] if i < len(calc_obj.ReactionY) else 0.0
                    rz = calc_obj.ReactionZ[i] if i < len(calc_obj.ReactionZ) else 0.0
                    mx = calc_obj.ReactionMX[i] if i < len(calc_obj.ReactionMX) else 0.0
                    my = calc_obj.ReactionMY[i] if i < len(calc_obj.ReactionMY) else 0.0
                    mz = calc_obj.ReactionMZ[i] if i < len(calc_obj.ReactionMZ) else 0.0
                    
                    # Format values
                    fx_formatted = float(rx) if not units_available else rx
                    fy_formatted = float(ry) if not units_available else ry
                    fz_formatted = float(rz) if not units_available else rz
                    mx_formatted = float(mx) if not units_available else mx
                    my_formatted = float(my) if not units_available else my
                    mz_formatted = float(mz) if not units_available else mz
                    
                    # Add row data
                    row_data = [
                        node_name,
                        f"{x_pos:.3f}",
                        f"{y_pos:.3f}",
                        f"{z_pos:.3f}",
                        f"{fx_formatted:.3f}",
                        f"{fy_formatted:.3f}",
                        f"{fz_formatted:.3f}",
                        f"{mx_formatted:.3f}",
                        f"{my_formatted:.3f}",
                        f"{mz_formatted:.3f}"
                    ]
                    
                    table_data.append(row_data)
                    
            # Fallback: Collect reaction data directly from model nodes
            elif model and hasattr(model, 'nodes') and model.nodes:
                # Collect reaction data for each supported node from model
                for node_name, node in model.nodes.items():
                    if self.is_node_supported(node):
                        # Get reaction values for current load combination - with safety checks
                        rx = node.RxnFX.get(load_combo, 0.0) if hasattr(node, 'RxnFX') else 0.0
                        ry = node.RxnFY.get(load_combo, 0.0) if hasattr(node, 'RxnFY') else 0.0
                        rz = node.RxnFZ.get(load_combo, 0.0) if hasattr(node, 'RxnFZ') else 0.0
                        mx = node.RxnMX.get(load_combo, 0.0) if hasattr(node, 'RxnMX') else 0.0
                        my = node.RxnMY.get(load_combo, 0.0) if hasattr(node, 'RxnMY') else 0.0
                        mz = node.RxnMZ.get(load_combo, 0.0) if hasattr(node, 'RxnMZ') else 0.0
                        
                        # Node coordinates
                        x_pos = node.X
                        y_pos = node.Y
                        z_pos = node.Z
                        
                        # Format values
                        fx_formatted = float(rx) if not units_available else rx
                        fy_formatted = float(ry) if not units_available else ry
                        fz_formatted = float(rz) if not units_available else rz
                        mx_formatted = float(mx) if not units_available else mx
                        my_formatted = float(my) if not units_available else my
                        mz_formatted = float(mz) if not units_available else mz
                        
                        # Add row data
                        row_data = [
                            node_name,
                            f"{x_pos:.3f}",
                            f"{y_pos:.3f}",
                            f"{z_pos:.3f}",
                            f"{fx_formatted:.3f}",
                            f"{fy_formatted:.3f}",
                            f"{fz_formatted:.3f}",
                            f"{mx_formatted:.3f}",
                            f"{my_formatted:.3f}",
                            f"{mz_formatted:.3f}"
                        ]
                        
                        table_data.append(row_data)
            
            # No valid reaction data found
            else:
                self.status_label.setText("No reaction data found - please run calculation first")
                self.status_label.setStyleSheet("color: #e74c3c;")
                return
            
            # Sort table data by node ID
            table_data.sort(key=lambda x: x[0])
            
            # Set up table
            self.table_widget.setRowCount(len(table_data))
            self.table_widget.setColumnCount(len(headers))
            self.table_widget.setHorizontalHeaderLabels(headers)
            
            # Populate table with data
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_data in enumerate(row_data):
                    item = QtWidgets.QTableWidgetItem(str(cell_data))
                    item.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEnabled)
                    self.table_widget.setItem(row_idx, col_idx, item)
            
            # Resize columns to fit content
            self.table_widget.resizeColumnsToContents()
            
            # Update status
            self.status_label.setText(f"Displaying {len(table_data)} supported nodes")
            self.status_label.setStyleSheet("color: #27ae60;")
            
        except Exception as e:
            FreeCAD.Console.PrintError(f"Error populating reaction table: {str(e)}\n")
            self.status_label.setText(f"Error: {str(e)}")
            self.status_label.setStyleSheet("color: #e74c3c;")
    
    def is_node_supported(self, node) -> bool:
        """Check if a node has any support conditions."""
        return (node.support_DX or node.support_DY or node.support_DZ or 
                node.support_RX or node.support_RY or node.support_RZ)
    
    def on_load_combination_changed(self, combo_name):
        """Handle load combination selection change."""
        self.populate_reaction_table()
    
    def export_to_format(self, format_type: str):
        """Export reaction results to specified format."""
        try:
            # Get file extension and filter based on format
            if format_type == "excel":
                file_filter = "Excel Files (*.xlsx);;All Files (*)"
                default_ext = ".xlsx"
            elif format_type == "word":
                file_filter = "Word Files (*.docx);;All Files (*)"
                default_ext = ".docx"
            else:  # csv
                file_filter = "CSV Files (*.csv);;All Files (*)"
                default_ext = ".csv"
            
            # Get file path from user
            file_path, selected_filter = QtWidgets.QFileDialog.getSaveFileName(
                self.form, 
                f"Export Reaction Results to {format_type.upper()}", 
                "", 
                file_filter
            )
            
            if not file_path:
                return
            
            # Ensure file has correct extension
            if not file_path.lower().endswith(default_ext):
                file_path += default_ext
            
            # Export based on format
            if format_type == "excel":
                success = self.export_to_excel(file_path)
            elif format_type == "word":
                success = self.export_to_word(file_path)
            else:  # csv
                success = self.export_to_csv(file_path)
            
            if success:
                self.status_label.setText(f"Reaction results exported to: {file_path}")
                self.status_label.setStyleSheet("color: #27ae60;")
                QtWidgets.QMessageBox.information(
                    self.form, 
                    "Export Successful", 
                    f"Reaction results successfully exported to:\n{file_path}"
                )
            else:
                self.status_label.setText("Error exporting reaction results")
                self.status_label.setStyleSheet("color: #e74c3c;")
                QtWidgets.QMessageBox.critical(
                    self.form, 
                    "Export Failed", 
                    "Failed to export reaction results. Check FreeCAD console for details."
                )
                
        except Exception as e:
            FreeCAD.Console.PrintError(f"Error exporting reaction results: {str(e)}\n")
            self.status_label.setText(f"Error: {str(e)}")
            self.status_label.setStyleSheet("color: #e74c3c;")
    
    def export_to_csv(self, file_path: str) -> bool:
        """Export reaction results to CSV file."""
        try:
            # Get current load combination
            load_combo = self.load_combo_dropdown.currentText()
            
            # Import units manager for formatting
            try:
                from .utils.units_manager import format_force, format_moment
                units_available = True
            except ImportError:
                units_available = False
                format_force = lambda x: f"{x:.3f}"
                format_moment = lambda x: f"{x:.3f}"
            
            # Get table data
            headers = []
            for col in range(self.table_widget.columnCount()):
                headers.append(self.table_widget.horizontalHeaderItem(col).text())
            
            rows = []
            for row in range(self.table_widget.rowCount()):
                row_data = []
                for col in range(self.table_widget.columnCount()):
                    item = self.table_widget.item(row, col)
                    row_data.append(item.text() if item else "")
                rows.append(row_data)
            
            # Write to CSV
            with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                
                # Write header
                writer.writerow([f"Reaction Results - Load Combination: {load_combo}"])
                writer.writerow([])  # Empty row
                writer.writerow(headers)
                
                # Write data
                writer.writerows(rows)
                
                # Write summary
                writer.writerow([])
                writer.writerow([f"Total Nodes: {len(rows)}"])
            
            return True
            
        except Exception as e:
            FreeCAD.Console.PrintError(f"Error exporting to CSV: {str(e)}\n")
            return False
    
    def export_to_excel(self, file_path: str) -> bool:
        """Export reaction results to Excel file."""
        try:
            # Try to import openpyxl for Excel export
            try:
                import openpyxl
            except ImportError:
                QtWidgets.QMessageBox.warning(
                    self.form,
                    "Export Failed",
                    "Excel export requires the 'openpyxl' Python package.\nPlease install it using pip: pip install openpyxl"
                )
                return False
            
            # Get current load combination
            load_combo = self.load_combo_dropdown.currentText()
            
            # Create workbook and worksheet
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Reaction Results"
            
            # Write header
            ws.append([f"Reaction Results - Load Combination: {load_combo}"])
            ws.append([])  # Empty row
            
            # Write column headers
            headers = []
            for col in range(self.table_widget.columnCount()):
                headers.append(self.table_widget.horizontalHeaderItem(col).text())
            ws.append(headers)
            
            # Write data
            for row in range(self.table_widget.rowCount()):
                row_data = []
                for col in range(self.table_widget.columnCount()):
                    item = self.table_widget.item(row, col)
                    row_data.append(item.text() if item else "")
                ws.append(row_data)
            
            # Write summary
            ws.append([])
            ws.append([f"Total Nodes: {self.table_widget.rowCount()}"])
            
            # Save workbook
            wb.save(file_path)
            return True
            
        except Exception as e:
            FreeCAD.Console.PrintError(f"Error exporting to Excel: {str(e)}\n")
            return False
    
    def export_to_word(self, file_path: str) -> bool:
        """Export reaction results to Word file."""
        try:
            # Try to import python-docx for Word export
            try:
                from docx import Document
            except ImportError:
                QtWidgets.QMessageBox.warning(
                    self.form,
                    "Export Failed",
                    "Word export requires the 'python-docx' Python package.\nPlease install it using pip: pip install python-docx"
                )
                return False
            
            # Get current load combination
            load_combo = self.load_combo_dropdown.currentText()
            
            # Create document
            doc = Document()
            doc.add_heading('Reaction Results', 0)
            doc.add_paragraph(f'Load Combination: {load_combo}')
            doc.add_paragraph()
            
            # Create table
            headers = []
            for col in range(self.table_widget.columnCount()):
                headers.append(self.table_widget.horizontalHeaderItem(col).text())
            
            num_cols = len(headers)
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
            
            # Add data
            for row in range(self.table_widget.rowCount()):
                row_cells = table.add_row().cells
                for col in range(self.table_widget.columnCount()):
                    item = self.table_widget.item(row, col)
                    row_cells[col].text = item.text() if item else ""
            
            # Add summary
            doc.add_paragraph()
            doc.add_paragraph(f'Total Nodes: {self.table_widget.rowCount()}')
            
            # Save document
            doc.save(file_path)
            return True
            
        except Exception as e:
            FreeCAD.Console.PrintError(f"Error exporting to Word: {str(e)}\n")
            return False
    
    def accept(self):
        """Accept changes and close panel."""
        FreeCADGui.Control.closeDialog()
    
    def reject(self):
        """Reject changes and close panel."""
        FreeCADGui.Control.closeDialog()
    
    def getStandardButtons(self):
        """Return standard buttons for the panel."""
        return int(QtWidgets.QDialogButtonBox.Ok)


# Command to open the reaction table panel
class CommandReactionTablePanel:
    """Command to open reaction results table panel."""
    
    def GetResources(self):
        return {
            'Pixmap': os.path.join(ICONPATH, "reaction.svg"),
            'MenuText': "Reaction Results Table", 
            'ToolTip': "Display reaction forces and moments in a table format"
        }
    
    def Activated(self):
        try:
            # Check if a reaction results object is selected
            selection = FreeCADGui.Selection.getSelection()
            reaction_obj = None
            calc_obj = None
            
            # First, try to find a reaction results object
            for obj in selection:
                if hasattr(obj, 'Proxy') and isinstance(obj.Proxy, ReactionResultsClass):
                    reaction_obj = obj
                    break
                # Check if we have a calc object selected
                elif hasattr(obj, 'Proxy') and hasattr(obj.Proxy, '__class__') and 'Calc' in obj.Proxy.__class__.__name__:
                    calc_obj = obj
            
            # If no reaction object found, but calc object is found or available in document
            if not reaction_obj:
                if not calc_obj:
                    # Try to find calc object in document
                    for obj in FreeCAD.ActiveDocument.Objects:
                        if hasattr(obj, 'Proxy') and hasattr(obj.Proxy, '__class__'):
                            if 'Calc' in obj.Proxy.__class__.__name__:
                                calc_obj = obj
                                break
                
                if not calc_obj:
                    QtWidgets.QMessageBox.warning(None, "Warning", 
                        "Please select or create a calculation object first.")
                    return
                
                # Check if analysis has been run
                if not hasattr(calc_obj, 'model') or not calc_obj.model:
                    # Try alternative attribute names
                    if hasattr(calc_obj, 'FEModel') and calc_obj.FEModel:
                        pass  # FEModel is available
                    else:
                        QtWidgets.QMessageBox.warning(None, "Warning", 
                            "Please run the structural analysis first.")
                        return
                
                # Create reaction results object
                from freecad.StructureTools.reaction_results import ReactionResults, ViewProviderReactionResults
                FreeCAD.Console.PrintMessage("Creating new reaction results object...\n")
                reaction_obj = FreeCAD.ActiveDocument.addObject("App::DocumentObjectGroupPython", "ReactionResults")
                ReactionResults(reaction_obj, calc_obj)
                ViewProviderReactionResults(reaction_obj.ViewObject)
                FreeCAD.ActiveDocument.recompute()
            
            # Now open the reaction table panel with the reaction_obj
            panel = ReactionTablePanel(reaction_obj)
            FreeCADGui.Control.showDialog(panel)
            
        except Exception as e:
            FreeCAD.Console.PrintError(f"Error opening reaction table panel: {str(e)}\n")
            QtWidgets.QMessageBox.critical(None, "Error", f"Failed to open reaction table panel: {str(e)}")
    
    def IsActive(self):
        return FreeCAD.ActiveDocument is not None


# Import ReactionResults class for isinstance check
try:
    from .reaction_results import ReactionResults as ReactionResultsClass
except ImportError:
    class ReactionResultsClass:
        pass


# Register the command
if hasattr(FreeCADGui, 'addCommand'):
    FreeCADGui.addCommand('ReactionTablePanel', CommandReactionTablePanel())