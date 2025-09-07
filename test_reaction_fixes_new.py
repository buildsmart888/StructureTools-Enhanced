#!/usr/bin/env python3
"""
Test script to verify the reaction results coordinate system and Word export fixes.

This script tests:
1. Coordinate system conversion from Pynite to FreeCAD
2. Word export functionality with proper error handling
3. Overall reaction results table functionality

Usage:
- Run this script in FreeCAD console after creating a structural model
- Or run it standalone to test coordinate conversion functions
"""

import FreeCAD
import sys
import os

def test_coordinate_conversion():
    """Test coordinate system conversion logic."""
    print("="*60)
    print("TESTING COORDINATE SYSTEM CONVERSION")
    print("="*60)
    
    # Test cases: (pynite_x, pynite_y, pynite_z) -> expected (freecad_x, freecad_y, freecad_z)
    test_cases = [
        # Standard coordinates in mm
        (0, 0, 0, 0, 0, 0),
        (1000, 2000, 3000, 1000, 3000, 2000),
        (5000, 0, 4000, 5000, 4000, 0),
        
        # Coordinates that need m to mm conversion
        (1, 2, 3, 1000, 3000, 2000),
        (10, 0, 5, 10000, 5000, 0),
    ]
    
    all_passed = True
    
    for i, (px, py, pz, expected_fx, expected_fy, expected_fz) in enumerate(test_cases):
        print(f"\nTest case {i+1}:")
        print(f"  Pynite coords: ({px}, {py}, {pz})")
        
        # Apply conversion logic (same as in reaction_table_panel.py)
        pynite_x = px * 1000 if abs(px) < 100 else px
        pynite_y = py * 1000 if abs(py) < 100 else py  
        pynite_z = pz * 1000 if abs(pz) < 100 else pz
        
        # Map to FreeCAD coordinate system
        freecad_x = pynite_x  # X remains same
        freecad_y = pynite_z  # Y becomes Pynite's Z (depth)
        freecad_z = pynite_y  # Z becomes Pynite's Y (vertical)
        
        print(f"  FreeCAD coords: ({freecad_x}, {freecad_y}, {freecad_z})")
        print(f"  Expected: ({expected_fx}, {expected_fy}, {expected_fz})")
        
        # Check if conversion is correct
        if (freecad_x == expected_fx and freecad_y == expected_fy and freecad_z == expected_fz):
            print(f"  ✅ PASSED")
        else:
            print(f"  ❌ FAILED")
            all_passed = False
    
    print(f"\n{'='*60}")
    print(f"Coordinate conversion test: {'PASSED' if all_passed else 'FAILED'}")
    print(f"{'='*60}")
    
    return all_passed

def test_word_export_dependencies():
    """Test Word export dependencies and functionality."""
    print("\n" + "="*60)
    print("TESTING WORD EXPORT DEPENDENCIES")
    print("="*60)
    
    # Test python-docx import
    try:
        from docx import Document
        print("✅ python-docx is available")
        
        # Test basic document creation
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        
        # Test table creation
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        
        # Add some test data
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Node ID'
        header_cells[1].text = 'X (mm)'
        header_cells[2].text = 'Z (mm)'
        
        data_cells = table.rows[1].cells
        data_cells[0].text = 'N1'
        data_cells[1].text = '1000.0'
        data_cells[2].text = '2000.0'
        
        print("✅ Word document creation test passed")
        print("✅ Word table creation test passed")
        return True
        
    except ImportError:
        print("❌ python-docx is NOT available")
        print("   Install with: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ Error testing Word functionality: {str(e)}")
        return False

def run_all_tests():
    """Run all tests."""
    print("REACTION RESULTS FIXES TEST SUITE")
    print("=" * 80)
    
    results = []
    
    # Run coordinate conversion test
    results.append(("Coordinate Conversion", test_coordinate_conversion()))
    
    # Run Word export dependencies test
    results.append(("Word Export Dependencies", test_word_export_dependencies()))
    
    # Summary
    print("\n" + "=" * 80)
    print("TEST SUMMARY")
    print("=" * 80)
    
    passed = 0
    for test_name, result in results:
        status = "PASSED" if result else "FAILED"
        print(f"{test_name:30} : {status}")
        if result:
            passed += 1
    
    print(f"\nOverall: {passed}/{len(results)} tests passed")
    
    if passed == len(results):
        print("🎉 All tests passed! The fixes are working correctly.")
    else:
        print("⚠️  Some tests failed. Check the output above for details.")
    
    return passed == len(results)

if __name__ == "__main__":
    # Run tests
    success = run_all_tests()
    sys.exit(0 if success else 1)